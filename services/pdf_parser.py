"""
文档解析服务 - 支持 PDF、Word、TXT
"""
import os
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
import io


class DocumentParser:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt']

    def extract_text(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.pdf':
            return self._extract_pdf(filepath)
        elif ext == '.docx':
            return self._extract_docx(filepath)
        elif ext == '.txt':
            return self._extract_txt(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _extract_pdf(self, filepath):
        text = ""
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                page_image = self._extract_pdf_page_image(page)
                if page_image:
                    ocr_text = self._ocr_image(page_image)
                    if ocr_text:
                        text += f"\n[图片OCR识别]\n{ocr_text}\n"
        return text

    def _extract_pdf_page_image(self, page):
        try:
            if '/XObject' in page['/Resources']:
                xobjects = page['/XObject'].get_object()
                for obj in xobjects.values():
                    if obj.get('/Subtype') == '/Image':
                        return None
        except Exception:
            pass
        return None

    def _ocr_image(self, image_data):
        try:
            image = Image.open(io.BytesIO(image_data))
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            return text.strip()
        except Exception as e:
            print(f"OCR 识别失败: {e}")
            return ""

    def _extract_docx(self, filepath):
        text = ""
        doc = Document(filepath)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    def _extract_txt(self, filepath):
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(filepath, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')

    def extract_pages(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)
        elif ext == '.docx':
            doc = Document(filepath)
            return len(doc.paragraphs)
        elif ext == '.txt':
            content = self._extract_txt(filepath)
            return len(content.split('\n'))
        return 0